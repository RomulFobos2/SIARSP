from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SRC_MD = ROOT / "docs" / "SIARSP_service_description.md"
OUT_DOCX = ROOT / "docs" / "SIARSP_service_description.docx"


def parse_markdown(md_text: str):
    lines = md_text.splitlines()
    blocks = []
    for line in lines:
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        elif line.strip() == "":
            blocks.append(("empty", ""))
        else:
            blocks.append(("p", line.strip()))
    return blocks


def build_docx(blocks):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    first_h1_done = False
    for kind, text in blocks:
        if kind == "h1" and not first_h1_done:
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = RGBColor(31, 73, 125)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            first_h1_done = True
        elif kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        elif kind == "h3":
            doc.add_heading(text, level=3)
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
        elif kind == "p":
            doc.add_paragraph(text)
        elif kind == "empty":
            doc.add_paragraph()

    return doc


def main():
    if not SRC_MD.exists():
        raise FileNotFoundError(f"Не найден исходный файл: {SRC_MD}")

    blocks = parse_markdown(SRC_MD.read_text(encoding="utf-8"))
    doc = build_docx(blocks)
    doc.save(OUT_DOCX)
    print(f"DOCX сформирован: {OUT_DOCX}")


if __name__ == "__main__":
    main()
